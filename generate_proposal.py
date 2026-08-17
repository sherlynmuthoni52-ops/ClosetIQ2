#!/usr/bin/env python3
"""
ClosetIQ Project Proposal Generator
Generates a filled Word document based on the template.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Set background color for a table cell."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_custom(doc, text, level=1):
    """Add a styled heading."""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    run.font.color.rgb = RGBColor(44, 62, 80)
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(18)
    elif level == 2:
        run.font.size = Pt(14)
    else:
        run.font.size = Pt(12)
    return heading

def add_paragraph_custom(doc, text, bold=False, italic=False, align=None):
    """Add a paragraph with optional formatting."""
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic
    return p

def add_bullet_paragraph(doc, text, level=0):
    """Add a bulleted paragraph."""
    p = doc.add_paragraph(text, style='List Bullet')
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    return p

# Create document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ============================
# TITLE PAGE
# ============================
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_para.add_run('CLOSETIQ')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(44, 62, 80)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Smart Wardrobe Inventory and Outfit Planner')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(52, 152, 219)

doc.add_paragraph()
doc.add_paragraph()

by_para = doc.add_paragraph()
by_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = by_para.add_run('By')
run.font.size = Pt(14)

doc.add_paragraph()

name_para = doc.add_paragraph()
name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = name_para.add_run('[Student Full Name]')
run.font.size = Pt(14)

reg_para = doc.add_paragraph()
reg_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = reg_para.add_run('Registration No: [XXXXXXXXX]')
run.font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph()

decl_para = doc.add_paragraph()
decl_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = decl_para.add_run('A Project Proposal Submitted in Partial Fulfillment of the Requirements')
run.font.size = Pt(12)

decl_para2 = doc.add_paragraph()
decl_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = decl_para2.add_run('for the Award of the Degree in Information Technology')
run.font.size = Pt(12)

doc.add_paragraph()

university_para = doc.add_paragraph()
university_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = university_para.add_run('[University Name]')
run.font.size = Pt(14)
run.font.bold = True

dept_para = doc.add_paragraph()
dept_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = dept_para.add_run('Department of Information Technology')
run.font.size = Pt(12)

doc.add_paragraph()

year_para = doc.add_paragraph()
year_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = year_para.add_run('2026')
run.font.size = Pt(14)
run.font.bold = True

doc.add_page_break()

# ============================
# DECLARATION AND APPROVAL
# ============================
add_heading_custom(doc, 'Declaration and Approval', level=1)

decl_text = (
    'I, [Student Full Name], hereby declare that this project proposal titled '
    '"ClosetIQ: Smart Wardrobe Inventory and Outfit Planner" is my original work '
    'and has not been submitted to any other institution or university for academic purposes. '
    'All sources of information have been duly acknowledged.'
)
add_paragraph_custom(doc, decl_text)

doc.add_paragraph()
add_paragraph_custom(doc, 'Student Signature: ______________________________    Date: ____________')

doc.add_paragraph()
doc.add_paragraph()

add_paragraph_custom(doc, 'Approved by:', bold=True)
doc.add_paragraph()
add_paragraph_custom(doc, 'Supervisor Name: ______________________________')
add_paragraph_custom(doc, 'Signature: ______________________________    Date: ____________')

doc.add_page_break()

# ============================
# ABSTRACT
# ============================
add_heading_custom(doc, 'Abstract', level=1)

abstract_text = (
    'ClosetIQ is a web-based Smart Wardrobe Inventory and Outfit Planner system designed '
    'to address the inefficiencies of manual wardrobe management. The system provides users '
    'with a digital platform to upload, categorize, and manage their clothing items under '
    'categories such as tops, bottoms, footwear, and accessories. Through integration with '
    'real-time weather data, ClosetIQ generates context-aware outfit suggestions tailored to '
    'current environmental conditions. The system employs a PHP-based backend, a MySQL relational '
    'database, and a responsive HTML/CSS/JavaScript frontend, developed and tested using XAMPP. '
    'By digitizing wardrobe inventories and delivering intelligent outfit recommendations, '
    'ClosetIQ simplifies daily outfit selection, promotes better clothing organization, and '
    'enhances the overall user dressing experience.'
)
add_paragraph_custom(doc, abstract_text)

doc.add_page_break()

# ============================
# TABLE OF CONTENTS
# ============================
add_heading_custom(doc, 'Table of Contents', level=1)

toc_items = [
    'Declaration and Approval\t2',
    'Abstract\t3',
    'Table of Contents\t4',
    'List of Tables\t6',
    'List of Figures\t7',
    'Chapter 1: Introduction\t8',
    '    1.1 Background of the System\t8',
    '    1.2 Problem Statement\t8',
    '    1.3 Objectives\t8',
    '    1.3.1 General Objective\t8',
    '    1.3.2 Specific Objectives\t8',
    '    1.4 Assumptions\t9',
    '    1.5 Scope and Limitations\t9',
    'Chapter 2: Literature Review\t10',
    '    2.1 Introduction\t10',
    '    2.2 Wardrobe Management and Digital Clothing Inventories\t10',
    '    2.3 Outfit Recommendation Systems\t10',
    '    2.4 Weather API Integration in Personal Applications\t10',
    '    2.5 Web Application Development Frameworks\t11',
    'Chapter 3: Development Methodology\t12',
    '    3.1 Development Methodology\t12',
    '    3.2 Functional and Non-Functional Requirements\t12',
    '    3.3 Design Tools\t13',
    '    3.4 Development Tools\t13',
    '    3.4.1 Database Development Tools\t13',
    '    3.4.2 Programming Tools\t13',
    '    3.5 Proposed System Modules\t14',
    '    3.6 Project Milestones and Schedule\t14',
    '    3.7 Project Budget\t15',
    'References\t16'
]

for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.tab_stops

doc.add_page_break()

# ============================
# LIST OF TABLES
# ============================
add_heading_custom(doc, 'List of Tables', level=1)
add_paragraph_custom(doc, 'Table 1: Project Milestones and Schedule')
add_paragraph_custom(doc, 'Table 2: Project Budget')

doc.add_page_break()

# ============================
# LIST OF FIGURES
# ============================
add_heading_custom(doc, 'List of Figures', level=1)
add_paragraph_custom(doc, 'Figure 1: Proposed System Architecture (to be developed during detailed design phase)')
add_paragraph_custom(doc, 'Figure 2: Use Case Diagram (to be developed during detailed design phase)')

doc.add_page_break()

# ============================
# CHAPTER 1
# ============================
doc.add_page_break()
add_heading_custom(doc, 'Chapter 1: Introduction', level=1)

# 1.1 Background
add_heading_custom(doc, '1.1 Background of the System', level=2)

bg_text1 = (
    'Fashion and personal presentation are an established part of daily and working life, '
    'and the number of garments an individual owns has grown steadily as clothing has become '
    'cheaper and more widely available. This growth has not been matched by an equivalent growth '
    'in the tools people use to keep track of what they own. Most individuals still rely on memory '
    'and periodic physical inspection of a wardrobe or closet to decide what to wear, a method that '
    'becomes progressively less reliable as a collection of clothing expands.'
)
add_paragraph_custom(doc, bg_text1)

bg_text2 = (
    'Working professionals are a group for whom this gap is particularly noticeable. Unlike casual '
    'dressing, professional dress is repeated on a near-daily basis, is often subject to workplace '
    'expectations of variety and neatness, and must be decided within a limited morning window before '
    'commuting to work. Over time, a professional wardrobe accumulates tops, bottoms, formal wear, '
    'and accessories acquired for different roles, seasons, and occasions, yet the same worker '
    'typically falls back on a small, familiar subset of these items simply because the rest are '
    'difficult to recall or to combine quickly.'
)
add_paragraph_custom(doc, bg_text2)

bg_text3 = (
    'Recent wardrobe-audit research supports this observation. A 2025 study of 156 adult wardrobes '
    'conducted jointly by KU Leuven and Utrecht University found that participants owned an average '
    'of 198 garments, of which approximately one-fifth had not been worn in the preceding twelve '
    'months despite being in good, reusable condition (Vermeyen et al., 2025). Comparable consumer '
    'research in other markets reports a similar pattern of clothing ownership outpacing clothing use '
    '(WRAP, 2022). These findings indicate that the challenge is not a shortage of clothing but a '
    'shortage of visibility into what is already owned and how it can be combined, which is precisely '
    'the kind of organisation and retrieval problem that an information system is designed to solve.'
)
add_paragraph_custom(doc, bg_text3)

bg_text4 = (
    'Advances in web and mobile technology have made it practical to address this gap directly. '
    'A lightweight web application can allow a user to photograph and log each garment once, tag it '
    'by category, colour, and season, and thereafter retrieve and combine that inventory in seconds '
    'rather than by physically searching a wardrobe. Because such a system already holds structured '
    'data about each item, it can also apply simple, transparent rules — for example, matching a top '
    'with a bottom of a compatible category, or narrowing suggestions using the day\'s weather forecast '
    'as a secondary, convenience-level filter — without requiring complex artificial intelligence or '
    'image-recognition technology that would be difficult to validate within an undergraduate project '
    'timeline.'
)
add_paragraph_custom(doc, bg_text4)

bg_text5 = (
    'It is this specific combination of circumstances — professionals with growing but poorly tracked '
    'wardrobes, a repeated daily decision made under time pressure, and the availability of '
    'straightforward web technologies capable of solving an information organisation problem — that '
    'motivates the proposed ClosetIQ system. The following sections define the problem more precisely, '
    'state the objectives that follow from it, and set out the boundaries within which the system will '
    'be developed.'
)
add_paragraph_custom(doc, bg_text5)

# 1.2 Problem Statement
add_heading_custom(doc, '1.2 Problem Statement', level=2)

prob_text1 = (
    'Working professionals who own a moderate-to-large personal wardrobe struggle to keep an accurate '
    'mental record of the clothing items they own and how those items can be combined into complete, '
    'weather-appropriate outfits. Because there is no structured record of the wardrobe, the same few '
    'combinations are worn repeatedly while other, perfectly usable items go unworn for long periods, '
    'and outfit selection each morning becomes a repeated, time-consuming task rather than a quick decision.'
)
add_paragraph_custom(doc, prob_text1)

prob_text2 = (
    'This pattern is not merely anecdotal. Wardrobe-audit research shows that a meaningful share of the '
    'clothing a person owns — on the order of one item in five — remains dormant for at least a year '
    'even though it is still fit for use (Vermeyen et al., 2025), and consumer surveys in other markets '
    'report the same imbalance between ownership and use (WRAP, 2022). For a working professional, the '
    'consequence of this imbalance is not simply clutter; it is a recurring cost paid every weekday '
    'morning in the time spent deciding what to wear, and a recurring financial cost in the form of '
    'duplicate purchases made because an existing, suitable item was forgotten or overlooked.'
)
add_paragraph_custom(doc, prob_text2)

prob_text3 = (
    'The problem is compounded by the fact that appropriate dress for a working professional is rarely '
    'a single fixed choice. It must additionally account for the day\'s weather, since a combination '
    'that is appropriate on a warm day may be unsuitable on a cold or rainy one. Without a system that '
    'keeps the wardrobe organised and cross-references it against current conditions, the professional '
    'must perform this weather check manually and separately from the process of choosing an outfit, '
    'adding friction to an already time-constrained routine.'
)
add_paragraph_custom(doc, prob_text3)

prob_text4 = (
    'Generic wardrobe or closet applications exist, but the review of existing solutions undertaken for '
    'this project (Chapter Two) shows that most focus narrowly on photo storage and manual tagging '
    'without providing a lightweight recommendation layer, and none combine categorised inventory '
    'management with real-time weather context in a way that is simple enough for a working professional '
    'to use in the few minutes available before leaving for work. There is therefore a need for a '
    'personal wardrobe information management system that lets a working professional record and '
    'categorise their clothing once, and thereafter retrieve suitable, weather-appropriate outfit '
    'combinations quickly, without the complexity or cost of a full commercial styling platform.'
)
add_paragraph_custom(doc, prob_text4)

prob_text5 = (
    'This project proposes ClosetIQ, a Personal Wardrobe Information Management and Outfit Planning '
    'System, to address this problem. The system\'s core purpose is inventory organisation and retrieval; '
    'weather-informed filtering is included as a secondary, rule-based enhancement to that core purpose '
    'rather than as an independent artificial-intelligence feature, in line with the scope appropriate '
    'to an undergraduate Information Systems project.'
)
add_paragraph_custom(doc, prob_text5)

# 1.3 Objectives
add_heading_custom(doc, '1.3 Objectives of the Project', level=2)

add_heading_custom(doc, '1.3.1 General Objective', level=3)
gen_obj = (
    'To design and develop a Personal Wardrobe Information Management and Outfit Planning System '
    '(ClosetIQ) that enables working professionals to record, categorise, and retrieve their clothing '
    'inventory, and to receive weather-informed outfit suggestions as a supporting feature of that inventory.'
)
add_paragraph_custom(doc, gen_obj)

add_heading_custom(doc, '1.3.2 Specific Objectives', level=3)

specific_objectives = [
    'To assess the challenges working professionals face in tracking their clothing inventory and selecting outfits within a limited daily routine.',
    'To review existing wardrobe management and outfit-planning applications and identify gaps relevant to working professionals.',
    'To design a relational database schema and system architecture for storing clothing inventory, user accounts, and outfit history.',
    'To develop a web-based system that allows a working professional to upload, categorise, and manage clothing items with image support.',
    'To implement a rule-based outfit suggestion feature, enhanced with weather-informed filtering, to support faster outfit decisions.',
    'To test and evaluate the developed system for functionality, usability, and its effect on wardrobe organisation and outfit-selection time.'
]

for obj in specific_objectives:
    add_bullet_paragraph(doc, obj)

# 1.4 Assumptions
add_heading_custom(doc, '1.4 Assumptions of the Project', level=2)

assumptions = [
    'Users will have access to a device with a modern web browser and a stable internet connection.',
    'Working professionals who take part in testing will be willing to photograph and upload their own clothing items for the purpose of evaluating the system.',
    'The OpenWeatherMap API will remain accessible during the development and testing period; the system\'s core inventory and categorisation features are designed to function correctly even if the weather data source is temporarily unavailable.',
    'Test users will interact with the system in English.',
    'The initial version of the system will be evaluated with a small group of individual users rather than an organisation-wide deployment.'
]

for assumption in assumptions:
    add_bullet_paragraph(doc, assumption)

# 1.5 Scope
add_heading_custom(doc, '1.5 Scope of the Project', level=2)

scope_intro = (
    'The project targets working professionals who wish to manage a personal wardrobe more efficiently '
    'and reduce the time spent selecting outfits. ClosetIQ shall be developed and delivered as a web-based '
    'application accessible through a standard browser. The scope of the current development phase includes:'
)
add_paragraph_custom(doc, scope_intro)

scope_items = [
    'User account registration, login, and secure authentication.',
    'Upload and management of clothing items with images and metadata such as category, colour, and season.',
    'Categorisation of clothing into tops, bottoms, footwear, and accessories, with filtering by category.',
    'Generation of outfit suggestions using rule-based logic, enhanced with real-time weather data obtained through the OpenWeatherMap API.',
    'A personal outfit history log and basic usage statistics for each user.'
]

for item in scope_items:
    add_bullet_paragraph(doc, item)

doc.add_paragraph()
add_paragraph_custom(doc, 'The following are explicitly outside the scope of the current development phase:', bold=True)

limitations = [
    'Social sharing or community features.',
    'Integration with e-commerce platforms for direct purchases.',
    'A native mobile application (iOS or Android).',
    'Artificial intelligence or machine-learning-based image recognition or style recommendation.',
    'Calendar-based outfit scheduling is treated as a possible future enhancement beyond the current scope and is not counted as a deliverable of this proposal.'
]

for limit in limitations:
    add_bullet_paragraph(doc, limit)

# References cited in chapter 1
add_heading_custom(doc, 'References Cited in This Chapter', level=2)

refs = [
    'Vermeyen, V., Alaerts, L., Worrell, E., & Van Acker, K. (2025). Behind closed doors: Examining the stock of clothing in individuals\' wardrobes. Journal of Circular Economy, 3(1). https://doi.org/10.55845/OQEE5977',
    'WRAP. (2022). Citizen insight: Clothing longevity and repair. Waste and Resources Action Programme.'
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    run = p.add_run(ref)
    run.font.size = Pt(10)

doc.add_page_break()

# ============================
# CHAPTER 2
# ============================
add_heading_custom(doc, 'Chapter 2: Literature Review', level=1)

add_heading_custom(doc, '2.1 Introduction', level=2)
intro_text = (
    'This chapter reviews existing literature and related systems relevant to wardrobe management, '
    'outfit planning, and the integration of external data sources in personal technology applications. '
    'The review draws from academic publications, technical documentation, and existing commercial systems '
    'to establish a theoretical foundation for ClosetIQ.'
)
add_paragraph_custom(doc, intro_text)

add_heading_custom(doc, '2.2 Wardrobe Management and Digital Clothing Inventories', level=2)
lit1 = (
    'Research by Niinimaki et al. (2020) highlights the phenomenon of wardrobe overconsumption, noting '
    'that a significant proportion of clothing items in an average wardrobe remain unused for extended '
    'periods. Digital wardrobe management tools have been proposed as a means of addressing this issue by '
    'increasing visibility and accessibility of owned garments. Studies in human-computer interaction '
    'suggest that visual cataloguing of clothing reduces redundant purchases and encourages more deliberate '
    'outfit assembly.'
)
add_paragraph_custom(doc, lit1)

lit2 = (
    'Commercial platforms such as Stylebook and Cladwell have demonstrated user interest in digital '
    'wardrobe inventories. However, these platforms are predominantly mobile-first and subscription-based, '
    'limiting accessibility for users who prefer lightweight, browser-based solutions.'
)
add_paragraph_custom(doc, lit2)

add_heading_custom(doc, '2.3 Outfit Recommendation Systems', level=2)
lit3 = (
    'Outfit recommendation is a complex task that involves compatibility matching across clothing '
    'attributes such as color, style, occasion, and season. Yu et al. (2019) proposed a compatibility '
    'learning framework that uses deep learning to assess outfit compatibility. While such machine learning '
    'approaches yield high accuracy, they require substantial training data and computational resources '
    'that are beyond the scope of an academic project at this level.'
)
add_paragraph_custom(doc, lit3)

lit4 = (
    'Rule-based recommendation systems provide a more accessible alternative for academic implementations. '
    'By defining outfit generation rules based on categories and external parameters such as weather, a '
    'functional recommendation engine can be developed using standard programming logic without reliance '
    'on machine learning infrastructure.'
)
add_paragraph_custom(doc, lit4)

add_heading_custom(doc, '2.4 Weather API Integration in Personal Applications', level=2)
lit5 = (
    'The integration of weather data into personal productivity and lifestyle applications is '
    'well-documented. OpenWeatherMap provides a freely accessible RESTful API that returns current weather '
    'conditions including temperature, humidity, and weather descriptions for any geographic location. '
    'Applications in the travel, fitness, and fashion domains have successfully leveraged this API to '
    'deliver context-aware user experiences.'
)
add_paragraph_custom(doc, lit5)

lit6 = (
    'For ClosetIQ, weather data will serve as a key input to the outfit recommendation engine, ensuring '
    'that suggested outfits are appropriate for the environmental conditions the user is currently '
    'experiencing.'
)
add_paragraph_custom(doc, lit6)

add_heading_custom(doc, '2.5 Web Application Development Frameworks', level=2)
lit7 = (
    'PHP, combined with MySQL, is one of the most widely used technology stacks for developing '
    'server-side web applications. Its maturity, extensive documentation, and compatibility with XAMPP '
    'make it an ideal choice for academic projects requiring a full-stack web solution. PHP supports '
    'prepared statements through PDO, ensuring secure database interactions, while MySQL provides a '
    'robust relational database management system suitable for structured data storage.'
)
add_paragraph_custom(doc, lit7)

lit8 = (
    'Frontend technologies including HTML5, CSS3, and JavaScript provide the interface layer for user '
    'interaction. Bootstrap 5 offers a responsive CSS framework that accelerates layout development and '
    'ensures cross-device compatibility. The OpenWeatherMap API provides a freely accessible RESTful '
    'interface for retrieving current weather conditions, which can be integrated server-side through '
    'PHP to protect API credentials.'
)
add_paragraph_custom(doc, lit8)

doc.add_page_break()

# ============================
# CHAPTER 3
# ============================
add_heading_custom(doc, 'Chapter 3: Development Methodology', level=1)

add_heading_custom(doc, '3.1 Development Methodology', level=2)

method_text1 = (
    'The Agile Software Development Methodology, specifically an iterative and incremental approach, '
    'will be adopted for the development of ClosetIQ. Agile was chosen because it supports flexible '
    'adaptation to changing requirements during development, encourages frequent testing and review, '
    'and allows for the delivery of working system components at the end of each iteration.'
)
add_paragraph_custom(doc, method_text1)

method_text2 = (
    'The development will be structured in short sprints of one to two weeks each. At the end of each '
    'sprint, a functional component of the system will be reviewed and refined before proceeding to the '
    'next. This approach reduces the risk of discovering fundamental design flaws at a late stage and '
    'ensures that the system is continuously tested throughout development.'
)
add_paragraph_custom(doc, method_text2)

method_text3 = (
    'The methodology will be applied as follows: requirements will be gathered and documented in the '
    'first iteration, system design artifacts will be produced in the second, followed by sequential '
    'development of the database, backend, and frontend components, with final integration and testing '
    'conducted in the last sprint.'
)
add_paragraph_custom(doc, method_text3)

# 3.2 Requirements
add_heading_custom(doc, '3.2 Functional and Non-Functional Requirements', level=2)

add_heading_custom(doc, '3.2.1 Functional Requirements', level=3)
functional_reqs = [
    'The system shall allow users to register and log in using a username and password.',
    'The system shall allow authenticated users to add, edit, and delete clothing items.',
    'The system shall support image upload for each clothing item.',
    'The system shall categorise clothing items into tops, bottoms, footwear, and accessories.',
    'The system shall retrieve current weather data based on the user\'s location input.',
    'The system shall generate outfit suggestions based on selected categories and weather conditions.',
    'The system shall maintain a log of previously generated outfits for each user.'
]
for req in functional_reqs:
    add_bullet_paragraph(doc, req)

add_heading_custom(doc, '3.2.2 Non-Functional Requirements', level=3)
non_func_reqs = [
    'Performance: The system shall load any page within 3 seconds under normal network conditions.',
    'Usability: The interface shall be intuitive and navigable without prior training.',
    'Security: User passwords shall be stored as hashed values using industry-standard hashing.',
    'Reliability: The system shall be available for use at all times except during scheduled maintenance.',
    'Scalability: The architecture shall support the addition of new clothing categories without major code refactoring.'
]
for req in non_func_reqs:
    add_bullet_paragraph(doc, req)

# 3.3 Design Tools
add_heading_custom(doc, '3.3 Design Tools', level=2)

design_text1 = (
    'The following design tools will be used during the system design phase:'
)
add_paragraph_custom(doc, design_text1)

design_tools = [
    ('Draw.io (diagrams.net)', 
     'This free, browser-based diagramming tool will be used to create Data Flow Diagrams (DFDs), '
     'Entity Relationship Diagrams (ERDs), and Use Case Diagrams. Draw.io was selected because it '
     'supports a wide range of diagram types, integrates with Google Drive, and requires no installation. '
     'It will be used to visually represent system architecture, data flows, and entity relationships '
     'before development begins.'),
    ('Figma (Free Tier)', 
     'Figma will be used to design wireframes and UI mockups for the web interface. Figma supports '
     'collaborative design and allows for rapid prototyping of user interface components. The tool will '
     'be used to finalize the layout of key screens such as the dashboard, clothing upload form, and '
     'outfit suggestion view prior to frontend coding.')
]

for tool_name, tool_desc in design_tools:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(tool_name + ': ')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(tool_desc)
    run.font.size = Pt(11)

# 3.4 Development Tools
add_heading_custom(doc, '3.4 Development Tools', level=2)

add_heading_custom(doc, '3.4.1 Database Development Tools', level=3)
db_tools = (
    'MySQL: MySQL is a widely-used open-source relational database management system. It will serve as '
    'the primary database for ClosetIQ, storing user account data, clothing inventory records, and outfit '
    'history. MySQL was chosen for its reliability, extensive documentation, and seamless integration with '
    'PHP. phpMyAdmin, included with XAMPP, will be used as a graphical interface for schema design, '
    'database creation, and data inspection during development.'
)
add_paragraph_custom(doc, db_tools)

add_heading_custom(doc, '3.4.2 Programming Tools', level=3)

prog_tools = [
    ('PHP (Core Language)', 
     'PHP is a mature, server-side scripting language designed for web development. It will be used to '
     'develop all server-side logic for ClosetIQ, including user authentication, routing, database '
     'interactions, and API integration. PHP was chosen because it is natively supported by XAMPP, '
     'has extensive community support, and allows rapid prototyping of web application features. '
     'All database queries will use PDO (PHP Data Objects) with prepared statements to ensure security.'),
    ('HTML5, CSS3, and JavaScript', 
     'These standard web technologies will form the frontend of ClosetIQ. HTML5 will structure the web '
     'pages with semantic elements, CSS3 will handle visual styling and responsive layout using Flexbox '
     'and Grid, and JavaScript will manage dynamic interactions such as category filtering, image '
     'previews, and AJAX-based weather data retrieval. Bootstrap 5 will be used as a CSS framework to '
     'accelerate responsive layout development.'),
    ('OpenWeatherMap API', 
     'This RESTful API will be integrated into the backend to retrieve current weather data based on '
     'user-specified location. The API returns JSON-formatted weather data including temperature, '
     'humidity, and condition descriptions, which will be used as inputs to the outfit recommendation '
     'engine. Server-side proxy implementation ensures the API key remains secure.'),
    ('Visual Studio Code', 
     'VS Code will serve as the primary Integrated Development Environment (IDE) for this project. '
     'Its lightweight nature, extensive extension marketplace (including PHP and MySQL support), and '
     'integrated terminal make it ideal for full-stack web development.'),
    ('Git and GitHub', 
     'Version control will be managed using Git, with the repository hosted on GitHub. This ensures '
     'that all code changes are tracked, and that a recovery point exists at every development milestone.')
]

for tool_name, tool_desc in prog_tools:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(tool_name + ': ')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(tool_desc)
    run.font.size = Pt(11)

# 3.5 Proposed System Modules
add_heading_custom(doc, '3.5 Proposed System Modules', level=2)

modules = [
    'User Authentication Module: Handles user registration, login, session management, and logout. Ensures that all wardrobe data is private and accessible only to the authenticated owner.',
    'Clothing Inventory Module: Enables users to add, view, edit, and delete clothing items. Each item record includes an image, name, category, color, size, and season suitability.',
    'Category Management Module: Manages the predefined clothing categories (tops, bottoms, footwear, accessories) and allows filtering of inventory by category.',
    'Weather Integration Module: Interfaces with the OpenWeatherMap API to retrieve current weather data based on a city name entered by the user. Returns temperature and condition data to inform outfit suggestions.',
    'Outfit Suggestion Module: Applies rule-based logic to generate outfit combinations from the user\'s inventory, factoring in selected categories and current weather conditions. For example, cold temperatures trigger suggestions of heavier tops and closed footwear.',
    'Outfit History Module: Records and displays previously generated outfit suggestions for each user, enabling reference and repeated use of successful combinations.',
    'Dashboard Module: Provides an overview of the user\'s wardrobe statistics, including total items per category, recent outfit suggestions, and current weather status.'
]

for module in modules:
    add_bullet_paragraph(doc, module)

# 3.6 Milestones
add_heading_custom(doc, '3.6 Project Milestones and Schedule', level=2)

# Add table
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
headers = ['Phase', 'Activity', 'Duration', 'Timeline']
for i, header in enumerate(headers):
    hdr_cells[i].text = header
    set_cell_shading(hdr_cells[i], '2c3e50')
    hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    hdr_cells[i].paragraphs[0].runs[0].font.bold = True

milestones = [
    ['Phase 1', 'Requirements Gathering & System Analysis', '2 weeks', 'Week 1–2'],
    ['Phase 2', 'System Design (ERD, DFD, UI Mockups)', '2 weeks', 'Week 3–4'],
    ['Phase 3', 'Database Development', '1 week', 'Week 5'],
    ['Phase 4', 'Backend Development (PHP/MySQL)', '3 weeks', 'Week 6–8'],
    ['Phase 5', 'Frontend Development (HTML/CSS/JS)', '2 weeks', 'Week 9–10'],
    ['Phase 6', 'Integration & Testing', '1 week', 'Week 11'],
    ['Phase 7', 'Documentation & Final Report', '1 week', 'Week 12']
]

for row_data in milestones:
    row_cells = table.add_row().cells
    for i, val in enumerate(row_data):
        row_cells[i].text = val

doc.add_paragraph()
note_text = (
    'Note: A detailed Gantt chart illustrating the above schedule will be included in the final '
    'project report, generated using Microsoft Project or an equivalent tool.'
)
add_paragraph_custom(doc, note_text, italic=True)

# 3.7 Budget
add_heading_custom(doc, '3.7 Project Budget', level=2)

# Add budget table
budget_table = doc.add_table(rows=1, cols=3)
budget_table.style = 'Table Grid'
budget_hdr = budget_table.rows[0].cells
budget_headers = ['Item', 'Description', 'Estimated Cost (KES)']
for i, header in enumerate(budget_headers):
    budget_hdr[i].text = header
    set_cell_shading(budget_hdr[i], '2c3e50')
    budget_hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    budget_hdr[i].paragraphs[0].runs[0].font.bold = True

budget_items = [
    ['Internet Access', 'Research, APIs, cloud services', '3,000'],
    ['Cloud Hosting', 'Deployment on PythonAnywhere or Render', '2,000'],
    ['OpenWeatherMap API', 'Weather data integration (free tier)', '0'],
    ['Development Software', 'VS Code, phpMyAdmin, Figma (free)', '0'],
    ['Printing & Binding', 'Final report submission copies', '1,500'],
    ['Stationery', 'Notebooks, pens for documentation', '500'],
    ['Contingency (10%)', 'Unforeseen expenses', '700'],
    ['Total', '', '7,700']
]

for row_data in budget_items:
    row_cells = budget_table.add_row().cells
    for i, val in enumerate(row_data):
        row_cells[i].text = val
        if row_data[0] == 'Total':
            row_cells[i].paragraphs[0].runs[0].font.bold = True

doc.add_paragraph()
budget_note = (
    'Note: The total estimated budget is KES 7,700. The majority of development tools and services '
    'will be accessed under free tiers, significantly reducing project costs.'
)
add_paragraph_custom(doc, budget_note, italic=True)

doc.add_page_break()

# ============================
# REFERENCES
# ============================
add_heading_custom(doc, 'References', level=1)

references = [
    'Grinberg, M. (2018). Flask web development: Developing web applications with Python (2nd ed.). O\'Reilly Media.',
    'Niinimaki, K., Peters, G., Dahlbo, H., Perry, P., Rissanen, T., & Gwilt, A. (2020). The environmental price of fast fashion. Nature Reviews Earth and Environment, 1(4), 189-200. https://doi.org/10.1038/s43017-020-0039-9',
    'OpenWeatherMap. (2024). Current weather data API documentation. Retrieved from https://openweathermap.org/current',
    'Pressman, R. S., & Maxim, B. R. (2019). Software engineering: A practitioner\'s approach (9th ed.). McGraw-Hill Education.',
    'Vermeyen, V., Alaerts, L., Worrell, E., & Van Acker, K. (2025). Behind closed doors: Examining the stock of clothing in individuals\' wardrobes. Journal of Circular Economy, 3(1). https://doi.org/10.55845/OQEE5977',
    'WRAP. (2022). Citizen insight: Clothing longevity and repair. Waste and Resources Action Programme.',
    'Yu, W., Zhang, H., He, X., Chen, X., Xiong, L., & Qin, Z. (2019). Aesthetic-based clothing recommendation. In Proceedings of the 2019 World Wide Web Conference (WWW \'19), 649-658. https://doi.org/10.1145/3308558.3313707'
]

for ref in references:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(ref)
    run.font.size = Pt(11)

# Save document
output_path = r'C:\Users\Admin\Desktop\IS PROJECT\ClosetIQ_Project_Proposal_Filled.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
